from __future__ import annotations

"""Module d'analyse de documents pour l'import."""

from typing import Any, Dict, List, Optional, Tuple, Union

import docx
import fitz  # PyMuPDF


def analyser_docx(
    file_stream,
) -> Tuple[List[Dict[str, str]], Optional[Dict[str, Any]]]:
    """Tente d'extraire le contenu structuré et un template de style simplifié d'un DOCX.

    Retourne ``(contenu_structure, styles)`` où ``styles`` est ``None`` si l'extraction
    de style échoue. ``contenu_structure`` est une liste de dictionnaires décrivant
    chaque bloc de contenu du document.
    """
    try:
        file_stream.seek(0)
        document = docx.Document(file_stream)

        contenu_structure: List[Dict[str, str]] = []
        for para in document.paragraphs:
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            block_type = "paragraph"
            if style_name.startswith("heading 1") or style_name.startswith("titre 1"):
                block_type = "heading_1"
            elif style_name.startswith("heading 2") or style_name.startswith("titre 2"):
                block_type = "heading_2"
            elif style_name.startswith("heading 3") or style_name.startswith("titre 3"):
                block_type = "heading_3"
            elif style_name.startswith("heading 4") or style_name.startswith("titre 4"):
                block_type = "heading_4"
            elif style_name.startswith("heading 5") or style_name.startswith("titre 5"):
                block_type = "heading_5"
            elif style_name.startswith("heading 6") or style_name.startswith("titre 6"):
                block_type = "heading_6"

            if para.text.strip():
                contenu_structure.append({"type": block_type, "text": para.text})

        styles: Optional[Dict[str, Any]] = None
        try:
            first_para = document.paragraphs[0] if document.paragraphs else None
            first_run = first_para.runs[0] if first_para and first_para.runs else None
            if first_run:
                font_name = first_run.font.name or "Calibri"
                font_size = first_run.font.size.pt if first_run.font.size else 11
                styles = {
                    "prompt": {
                        "font_name": font_name,
                        "font_size": font_size,
                        "is_bold": bool(first_run.bold),
                    },
                    "response": {
                        "font_name": font_name,
                        "font_size": font_size,
                        "is_bold": False,
                    },
                }
        except Exception:
            styles = None

        return contenu_structure, styles
    except Exception:
        try:
            file_stream.seek(0)
            document = docx.Document(file_stream)
            contenu_structure = []
            for para in document.paragraphs:
                if para.text.strip():
                    contenu_structure.append({"type": "paragraph", "text": para.text})
        except Exception:
            contenu_structure = []
        return contenu_structure, None


def analyser_pdf(file_stream) -> Tuple[str, None]:
    """Extrait le contenu textuel brut d'un PDF."""
    file_stream.seek(0)
    with fitz.open(stream=file_stream.read(), filetype="pdf") as doc:
        full_text = "".join(page.get_text() for page in doc)
    return full_text, None


def analyser_document(
    fichier,
) -> Tuple[Union[str, List[Dict[str, str]]], Optional[Dict[str, Any]]]:
    """Analyse un fichier importé et choisit la méthode appropriée."""
    filename = fichier.name.lower()
    if filename.endswith(".docx"):
        return analyser_docx(fichier)
    if filename.endswith(".pdf"):
        return analyser_pdf(fichier)
    return "", None
